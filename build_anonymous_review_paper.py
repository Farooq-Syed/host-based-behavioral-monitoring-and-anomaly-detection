from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).with_name("HBBMAAD_anonymous_methodological_review.docx")

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(85, 85, 85)
LIGHT = "F2F4F7"

def set_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW") or OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    if tblW.getparent() is None: tblPr.append(tblW)
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa"); tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths): col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcW = cell._tc.tcPr.tcW
            tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def cell_text(cell, text, bold=False, center=False, size=9):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); set_font(r, size=size, bold=bold)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, 9, color=MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p

def add_para(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.13
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); set_font(r, 10.5, italic=italic)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text); set_font(r, 10.2)
    return p

def add_results_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        shade(table.rows[0].cells[i], LIGHT); cell_text(table.rows[0].cells[i], h, bold=True, center=i>0, size=8.5)
    trPr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row): cell_text(cells[i], value, center=i>0, size=8.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.82)
    sec.left_margin = sec.right_margin = Inches(0.82)
    sec.header_distance = sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal.font.size = Pt(10.5)
    for name, size, color, before, after in [("Heading 1", 15, BLUE, 14, 6), ("Heading 2", 12, INK, 10, 4)]:
        st = doc.styles[name]; st.font.name = "Calibri"; st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st.font.size = Pt(size); st.font.color.rgb = color; st.font.bold = True
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)

    hdr = sec.header.paragraphs[0]
    hdr.text = "Anonymous manuscript for methodological review"
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in hdr.runs: set_font(r, 8.5, italic=True, color=MUTED)
    add_page_number(sec.footer.paragraphs[0])

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28); p.paragraph_format.space_after = Pt(8)
    r = p.add_run("When Ransomware Detectors Meet Unseen Families")
    set_font(r, 22, bold=True, color=INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("A strict Sysmon evaluation of rules, anomaly detectors, supervised learning, and ensembles")
    set_font(r, 12, italic=True, color=MUTED)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Anonymous submission for independent methodological feedback")
    set_font(r, 10, color=MUTED)

    add_heading(doc, "Abstract")
    add_para(doc, "Endpoint ransomware detection is often reported on random train/test splits that mix behavior from the same malware family across folds. This manuscript compares rule-based detection, Isolation Forest, Local Outlier Factor, a supervised Random Forest, and two ensembles on public RADAR Sysmon telemetry. The evaluation holds out each ransomware family in turn and tests against the unseen family plus a fixed random 20% goodware pool. On seven held-out families, Random Forest achieved macro F1 0.400 (95% CI +/-0.13) and AUC 0.809; Isolation Forest and LOF achieved F1 0.324 and 0.195. However, performance varied sharply by family, from F1 0.514 to 0.200. A 1% FPR cutoff was selected solely on inner validation: macro-average achieved test FPR was 0.005 (range 0.000-0.011) and recall was 0.166. The result supports a narrow conclusion: conventional supervised detection outperforms the tested unsupervised baselines on these RADAR windows, but performance is sensitive to unseen-family shift. It does not establish host/session-disjoint or deployment-level generalization.")
    add_heading(doc, "Keywords", 2)
    add_para(doc, "ransomware detection; Sysmon; malware-family generalization; anomaly detection; evaluation methodology; false-positive-rate calibration")

    add_heading(doc, "1. Introduction")
    add_para(doc, "Ransomware detectors are commonly evaluated using endpoint or event telemetry from a finite set of observed attack families. A high score under random cross-validation can be useful for measuring in-distribution classification, but it does not answer the more operational question: does the detector retain utility when the ransomware family changes? This distinction matters because malware families differ in their staging, registry behavior, process creation, and file operations.")
    add_para(doc, "This study asks two limited questions. First, how do simple rules, unsupervised anomaly detectors, a supervised classifier, and small ensembles compare under a common feature representation? Second, how much does performance vary when the attack family is absent from training? The aim is not to propose a new learning algorithm. It is to make a conventional comparison auditable and to state precisely what its dataset and split can support.")

    add_heading(doc, "2. Data and Windowed Telemetry")
    add_para(doc, "The study uses RADAR, a public dataset of Sysmon logs from seven ransomware families (Akira, BlackBasta, CyberVolk, LockBit, Lynx, Medusa, and Meow) and goodware activity. Raw events are aggregated into five-minute windows. The reproducible subset contains 4,300 windows: 3,740 benign and 560 ransomware windows. Each window contains ten Sysmon-recoverable event-count features: process creation, file creation/deletion, rename-like activity, network connections, registry events, module loads, unsigned process events, suspicious extension events, shadow-copy events, and total event count.")
    add_para(doc, "The representation intentionally uses what the source Sysmon logs can recover. It does not claim CPU, memory, raw file entropy, or complete enterprise context. Goodware is represented by a single run; therefore, the benign test pool is a random 20% split, not a separate session-disjoint goodware population.")

    add_heading(doc, "3. Detectors and Comparators")
    add_para(doc, "The rule baseline sums positive indicators for rename-like activity, suspicious extensions, shadow-copy events, unsigned process activity, deletion volume, and registry volume; three indicators trigger an alert. Isolation Forest and Local Outlier Factor are trained on standardized numerical features. The supervised baseline is a class-balanced Random Forest. An unweighted vote combines Isolation Forest, LOF, and Random Forest; a weighted ensemble combines four core methods using training-fold precision weights. The progression detector is retained as a deliberately limited comparator: it requires a recon-to-tamper-to-encrypt sequence on the same host/time stream.")
    add_para(doc, "The goal of including these conventional methods is diagnostic. They expose different operating points: high-recall anomaly detectors, a supervised classifier whose decision boundary is learned from known families, and conservative combinations that may trade recall for precision.")

    add_heading(doc, "4. Evaluation Protocol")
    add_para(doc, "The primary evaluation is leave-one-family-out. For each ransomware family, all windows from that family are withheld for testing. The test set also includes the same deterministic 20% random sample of goodware windows; the remaining goodware and all other ransomware families form the training fold. Thus, the protocol is a cross-family attack test with a held-out random benign pool. It is not a host/session-disjoint deployment test.")
    add_para(doc, "Isolation Forest and LOF contamination are selected only within the training fold using inner validation. Random Forest and ensemble classification thresholds remain fixed at 0.5. For the operational recall@1%FPR figure, an additional Random Forest threshold is selected on an inner validation split: among thresholds with validation FPR no greater than 1%, the one with maximal validation recall is used once on the untouched test family. The achieved test FPR is reported separately because it can vary under family shift.")
    add_bullet(doc, "No held-out ransomware-family labels are used for contamination tuning or threshold selection.")
    add_bullet(doc, "Metrics are macro-averaged across the seven held-out families with t-based 95% confidence intervals.")
    add_bullet(doc, "Random cross-validation results are treated as development-era in-distribution reference results, not as the publication headline.")

    add_heading(doc, "5. Results")
    add_para(doc, "Under the strict family protocol, the supervised Random Forest was the strongest individual learned detector. The unweighted vote had slightly higher macro F1 but lower AUC is not defined for its binary vote. The weighted ensemble was highly precise but recalled few attacks. The rules were intentionally conservative and the progression detector did not flag RADAR trajectories because the available per-host groups were too short to establish the required three-stage sequence.")
    add_results_table(doc,
        ["Comparator", "F1 (95% CI)", "Precision", "Recall", "AUC", "Recall @ 1% FPR"],
        [["Random Forest", "0.400 (+/-0.13)", "0.347", "0.531", "0.809", "0.166"],
         ["Unweighted vote", "0.407 (+/-0.13)", "0.388", "0.462", "-", "-"],
         ["Isolation Forest", "0.324 (+/-0.11)", "0.256", "0.490", "-", "-"],
         ["Weighted ensemble", "0.208 (+/-0.06)", "0.967", "0.118", "-", "-"],
         ["LOF", "0.195 (+/-0.10)", "0.152", "0.296", "-", "-"],
         ["Rule baseline", "0.057 (+/-0.08)", "0.714", "0.030", "-", "-"],
         ["Progression", "0.000", "-", "-", "-", "-"]],
        [2050, 1450, 1050, 950, 850, 2200])
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Table 1. Macro-average strict leave-one-family-out results on 4,300 RADAR Sysmon windows.")
    set_font(r, 8.5, italic=True, color=MUTED)
    add_para(doc, "The difference across unseen families is material. The well-sampled Akira and BlackBasta families reached F1 about 0.51, while CyberVolk and Meow reached 0.20. Small families also produce wide uncertainty. This evidence does not show that one detector will work reliably for a novel ransomware family in a live enterprise; it shows that the apparent advantage observed under random cross-validation weakens under a stricter family split.")
    add_results_table(doc,
        ["Held-out family", "Attack windows", "RF F1", "RF AUC", "Recall @ 1%", "Achieved test FPR"],
        [["Akira", "117", "0.514", "0.808", "0.179", "0.005"],
         ["BlackBasta", "128", "0.511", "0.829", "0.141", "0.000"],
         ["LockBit", "126", "0.480", "0.796", "0.198", "0.005"],
         ["Medusa", "80", "0.464", "0.806", "0.225", "0.005"],
         ["Lynx", "56", "0.429", "0.861", "0.268", "0.011"],
         ["CyberVolk", "37", "0.200", "0.741", "0.027", "0.001"],
         ["Meow", "16", "0.200", "0.820", "0.125", "0.009"]],
        [1750, 1200, 1000, 1000, 1550, 2860])
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Table 2. Per-family Random Forest outcomes. The 1% constraint is imposed on inner validation; test FPR is measured, not constrained, on the unseen family.")
    set_font(r, 8.5, italic=True, color=MUTED)

    add_heading(doc, "6. Interpretation")
    add_para(doc, "The strict results support three observations. First, labelled supervised learning remains useful within the RADAR feature representation, outperforming the individual unsupervised baselines in macro F1 and yielding an AUC of 0.809. Second, the weak rule baseline and the precision-heavy weighted ensemble illustrate why a high precision number should not be interpreted alone: both leave substantial attack activity unflagged. Third, the family-to-family variation is the main finding. Random-CV F1 of 0.515 is in-distribution optimistic compared with strict macro F1 of 0.400, and the family-specific scores make the average less reassuring than a single headline would suggest.")
    add_para(doc, "The calibration result should likewise be read narrowly. The threshold was selected without test labels and had macro test FPR 0.005, but one held-out family reached 0.011. This is expected under distribution shift and demonstrates why per-family achieved FPR should accompany a recall-at-budget statement.")

    add_heading(doc, "7. Limitations")
    add_bullet(doc, "RADAR logs are controlled laboratory executions of real ransomware and goodware, not longitudinal live-enterprise telemetry.")
    add_bullet(doc, "Goodware is a single run, so the benign component is randomly held out rather than session-disjoint.")
    add_bullet(doc, "Several ransomware families contain few windows; their point estimates have substantial uncertainty.")
    add_bullet(doc, "The progression comparator is an omission on this source because the available traces do not support a full staged trajectory; it is not evidence that temporal detection is ineffective.")
    add_bullet(doc, "Precision weights for the weighted ensemble are fitted on training predictions and should be treated as exploratory until evaluated with out-of-fold calibration.")

    add_heading(doc, "8. Reproducibility and Next Steps")
    add_para(doc, "The repository includes a raw-log adapter, a manifest of the RADAR archive, fixed seeds, strict-evaluation scripts, output JSON, dependency lockfile, and automated regression tests. The next empirical step is to evaluate on a longer-horizon endpoint corpus with genuine host/session grouping and to add RADAR's drift and imbalance settings. A second methodological improvement would calibrate ensemble weights through an inner validation or out-of-fold procedure rather than training-set precision.") 

    add_heading(doc, "8.1 Repository, data, and reproduction", 2)
    add_para(doc, "Repository (anonymized for review): github.com/Farooq-Syed/host-based-behavioral-monitoring-and-anomaly-detection. Entry points: radar_strict_eval.py (strict leave-one-family-out evaluation), sysmon_adapter.py (raw-log to windowed telemetry), real_data_eval.py (development random-CV), plus scripts/reproduce_radar.py and scripts/download_real_data.py. Frozen preprocessing (window size, feature list, seeds, label rule) is recorded in radar_manifest.json and requirements-lock.txt.")
    add_para(doc, "Data attribution and license. RADAR (J. Ispahany, M. R. Islam, M. A. Khan, M. Z. Islam, \u201cRADAR: a realistic dataset for advancing ransomware detection\u201d), Zenodo doi:10.5281/zenodo.14564541, CC BY 4.0. It is a public research dataset; the code and this manuscript are released under the repository's Non-Commercial Personal-Use License.")
    add_para(doc, "Reproduction commands. python -m pip install -r requirements-lock.txt; python scripts/reproduce_radar.py; python radar_strict_eval.py --input data/radar_real_windows_with_family.csv --label-column label --split family --family-column family --metrics-output results/radar_strict_family_eval.json; python real_data_eval.py --input data/radar_real_windows.csv --label-column label --contamination 0.13; python -m pytest -q")

    add_heading(doc, "8.2 AI-use disclosure", 2)
    add_para(doc, "AI coding assistance was used during implementation and drafting. The author directed the research question, the benchmark evaluation protocol, the strict split and recall@FPR calibration design, the interpretation of the negative family-shift result, and reviewed and verified the final code and manuscript claims. AI assistance did not set the research direction or the claims.")

    add_heading(doc, "9. Conclusion")
    add_para(doc, "This work does not introduce a new ransomware detector. Its contribution is an auditable comparison showing that a familiar supervised model outperforms the tested unsupervised methods on RADAR Sysmon windows while weakening under unseen ransomware-family shift. The strict protocol and explicit limits are more important than the absolute score: the study supports cross-family evidence on a public corpus, not a deployment-level generalization claim.")

    add_heading(doc, "References")
    for ref in [
        "[1] J. Ispahany, M. R. Islam, M. A. Khan, and M. Z. Islam. RADAR: A realistic dataset for advancing ransomware detection. Zenodo, doi:10.5281/zenodo.14564541.",
        "[2] F. T. Liu, K. M. Ting, and Z.-H. Zhou. Isolation Forest. 2008 IEEE International Conference on Data Mining, 2008.",
        "[3] M. M. Breunig et al. LOF: Identifying Density-Based Local Outliers. SIGMOD Record, 2000.",
        "[4] L. Breiman. Random Forests. Machine Learning, 2001.",
        "[5] M. A. F. da Costa et al. A systematic review of ransomware detection approaches. Computers & Security, 2023."
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run(ref); set_font(r, 9.2)

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = "Anonymous Manuscript"
    doc.core_properties.subject = "Methodological review"
    doc.core_properties.comments = ""
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__": main()
